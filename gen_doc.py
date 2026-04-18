#!/usr/bin/env python3
"""Generate Flow Control project user manual as Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

code_style = doc.styles.add_style('CodeBlock', 1)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.left_indent = Cm(0.5)
shading = code_style.paragraph_format._element.get_or_add_pPr()
shd = shading.makeelement(qn('w:shd'), {
    qn('w:val'): 'clear',
    qn('w:color'): 'auto',
    qn('w:fill'): 'F0F0F0'
})
shading.append(shd)


def add_code(text):
    for line in text.strip().split('\n'):
        doc.add_paragraph(line, style='CodeBlock')


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


# ====================================================================
#  TITLE PAGE
# ====================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Flow Control')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('SystemVerilog Traffic Scheduling & Shaping Engine\nUser Manual')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Version 1.0\n2026-04-18')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ====================================================================
#  TABLE OF CONTENTS
# ====================================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('1. Project Overview', 1),
    ('2. Project Structure', 1),
    ('3. Quick Start', 1),
    ('4. Core Concepts', 1),
    ('  4.1 Traffic Models', 2),
    ('  4.2 Traffic Queue', 2),
    ('  4.3 Scheduler', 2),
    ('  4.4 Port Shaper', 2),
    ('  4.5 Traffic Monitor', 2),
    ('  4.6 Flow Controller', 2),
    ('5. API Reference', 1),
    ('  5.1 Enumerations', 2),
    ('  5.2 Structures', 2),
    ('  5.3 flow_controller Methods', 2),
    ('6. Usage Examples', 1),
    ('  6.1 Basic Rate-Limited Traffic', 2),
    ('  6.2 Token Bucket with Burst', 2),
    ('  6.3 Multi-Queue WRR Scheduling', 2),
    ('  6.4 Strict Priority Scheduling', 2),
    ('  6.5 Mixed SP+WRR Scheduling', 2),
    ('  6.6 Burst Traffic Pattern', 2),
    ('  6.7 Random Traffic Model', 2),
    ('  6.8 Step Rate Profile', 2),
    ('  6.9 Custom Traffic Model', 2),
    ('  6.10 Statistics & Monitoring', 2),
    ('  6.11 Complete Test Bench', 2),
    ('7. Test Suite', 1),
    ('8. Building & Running', 1),
]
for item, level in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if level == 1:
        p.runs[0].bold = True

doc.add_page_break()

# ====================================================================
#  1. PROJECT OVERVIEW
# ====================================================================
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'Flow Control is a SystemVerilog-based traffic scheduling and shaping engine '
    'designed for simulation with Synopsys VCS. It provides a complete framework for '
    'generating, scheduling, shaping, and monitoring network traffic at the packet level.'
)
doc.add_paragraph('Key features:')
for f in [
    'Five built-in traffic models: constant rate, token bucket, burst, random, and step profile',
    'Three scheduling algorithms: Strict Priority (SP), Weighted Round Robin (WRR), and Mixed SP+WRR',
    'Port-level token bucket rate limiter',
    'Real-time statistics collection with windowed rate monitoring',
    'Extensible via custom traffic models',
    'Comprehensive test suite with 89 test cases',
]:
    doc.add_paragraph(f, style='List Bullet')

# ====================================================================
#  2. PROJECT STRUCTURE
# ====================================================================
doc.add_heading('2. Project Structure', level=1)
add_code('''flow_control/
+-- Makefile                        # Build system
+-- filelist.f                      # VCS compilation file list
+-- src/
|   +-- common/
|   |   +-- flow_defines.sv         # Type definitions and enums
|   +-- models/
|   |   +-- traffic_model_base.sv   # Abstract base class
|   |   +-- rate_model.sv           # Constant rate model
|   |   +-- token_bucket_model.sv   # Rate + burst model
|   |   +-- burst_model.sv          # Bursty traffic model
|   |   +-- random_model.sv         # Randomized inter-packet gaps
|   |   +-- step_model.sv           # Time-segmented rate profiles
|   +-- core/
|   |   +-- traffic_queue.sv        # Per-queue FIFO container
|   |   +-- scheduler.sv            # Queue selection (SP/WRR/Mixed)
|   |   +-- port_shaper.sv          # Token bucket port rate limiter
|   |   +-- flow_controller.sv      # Top-level orchestrator
|   +-- monitor/
|       +-- traffic_monitor.sv      # Statistics & windowed rate check
+-- test/
    +-- test_traffic_models.sv      # Traffic model unit tests
    +-- test_queue.sv               # Queue FIFO tests
    +-- test_scheduler.sv           # Scheduling algorithm tests
    +-- test_port_shaper.sv         # Token bucket tests
    +-- test_monitor.sv             # Statistics collection tests
    +-- test_flow_controller.sv     # Integration tests''')

doc.add_heading('Architecture Diagram', level=2)
add_code('''Input Packets --> Queue[N] + Model --> Scheduler --> Port Shaper --> Output
                     |                                            |
               Traffic Monitor <-------- record() --------  Statistics

Per-Queue Control:
  - Independent traffic model (rate, burst, random, step, or custom)
  - Priority or weight-based scheduling
  - Real-time statistics collection

Global Control:
  - Total port rate limit via token bucket
  - Multiple scheduling modes (SP, WRR, hybrid)
  - Configurable duration or packet count limits''')

# ====================================================================
#  3. QUICK START
# ====================================================================
doc.add_heading('3. Quick Start', level=1)
doc.add_paragraph(
    'The following minimal example creates a flow controller with one queue at 1 Gbps, '
    'sends 100 packets of 1000 bytes each, and prints a statistics report.'
)
add_code('''program quick_start;
    `include "core/flow_controller.sv"

    initial begin
        flow_controller fc = new();
        byte unsigned pkt[$];
        int qid;
        byte unsigned out_data[$];
        realtime send_time;

        // 1. Configure
        fc.set_port_rate(10000.0);          // 10 Gbps port
        fc.set_duration(100us);             // Run for 100 us

        // 2. Add queue with rate model at 1 Gbps
        fc.add_queue(0, MODEL_RATE);
        fc.set_queue_rate(0, 1000.0);       // 1 Gbps

        // 3. Generate packets
        pkt = {};
        for (int i = 0; i < 1000; i++) pkt.push_back(i % 256);
        for (int i = 0; i < 100; i++)
            fc.push_packet(0, pkt);

        // 4. Schedule and send
        fc.start_schedule();
        while (fc.get_next_scheduled(qid, out_data, send_time))
            ;  // Process out_data if needed

        // 5. Print report
        fc.report();
    end
endprogram''')

doc.add_paragraph('Compile and run:')
add_code('''# Set VCS environment
export VCS_HOME=/opt/synopsys/vcs/Q-2020.03-SP2-7
export PATH=$VCS_HOME/bin:$PATH

# Compile and simulate
vcs -full64 -sverilog -timescale=1ns/1ps \\
    -f filelist.f +incdir+src quick_start.sv -o simv
./simv''')

# ====================================================================
#  4. CORE CONCEPTS
# ====================================================================
doc.add_heading('4. Core Concepts', level=1)

doc.add_heading('4.1 Traffic Models', level=2)
doc.add_paragraph(
    'Traffic models control the inter-packet gap (IPG) for each queue. '
    'All models inherit from the abstract base class traffic_model_base and implement '
    'the get_interval(pkt_size, current_time) method, which returns the time to wait '
    'before the next packet.'
)

add_table(
    ['Model', 'Type Enum', 'Description', 'Key Parameters'],
    [
        ['rate_model', 'MODEL_RATE', 'Constant rate, evenly-spaced packets', 'rate_mbps'],
        ['token_bucket_model', 'MODEL_TOKEN_BUCKET', 'Rate with burst allowance', 'rate_mbps, burst_size_bytes'],
        ['burst_model', 'MODEL_BURST', 'Send N packets, then pause', 'send_count, pause_time'],
        ['random_model', 'MODEL_RANDOM', 'Stochastic IPG', 'avg_rate_mbps, dist_type'],
        ['step_model', 'MODEL_STEP', 'Time-segmented rate profile', 'steps[$], loop_enable'],
        ['(custom)', 'MODEL_CUSTOM', 'User-defined model', '(user-defined)'],
    ]
)

doc.add_paragraph()
doc.add_heading('rate_model', level=3)
doc.add_paragraph(
    'Generates evenly-spaced packets at a constant rate. '
    'The inter-packet gap is calculated as: gap = (pkt_size * 8) / (rate_mbps * 1e6) seconds.'
)

doc.add_heading('token_bucket_model', level=3)
doc.add_paragraph(
    'Implements a token bucket algorithm. Tokens refill at the configured rate, '
    'capped at burst_size_bytes. If sufficient tokens are available, the packet '
    'is sent immediately (gap = 0). Otherwise, the gap equals the time needed to '
    'refill enough tokens for the packet.'
)

doc.add_heading('burst_model', level=3)
doc.add_paragraph(
    'Sends send_count packets with zero gap (back-to-back burst), then waits '
    'pause_time before the next burst. Useful for simulating bursty traffic patterns.'
)

doc.add_heading('random_model', level=3)
doc.add_paragraph('Generates random inter-packet gaps with three distribution options:')
add_table(
    ['Distribution', 'Enum', 'Description'],
    [
        ['Uniform', 'DIST_UNIFORM', 'Gap uniformly distributed in [0.5*mean, 1.5*mean]'],
        ['Poisson', 'DIST_POISSON', 'Exponential distribution (Poisson process)'],
        ['Normal', 'DIST_NORMAL', 'Gaussian with sigma = 0.2 * mean'],
    ]
)

doc.add_paragraph()
doc.add_heading('step_model', level=3)
doc.add_paragraph(
    'Supports time-segmented rate profiles. Each step has a duration and rate. '
    'Steps are traversed sequentially. When loop_enable is set, the profile repeats.'
)

doc.add_heading('4.2 Traffic Queue', level=2)
doc.add_paragraph('traffic_queue is a per-queue FIFO container. Each queue has:')
for item in [
    'queue_id: Unique identifier',
    'model: Associated traffic model (determines IPG)',
    'prio: Priority level (higher = more priority, used by SP scheduler)',
    'weight: Weight (used by WRR scheduler, higher = more bandwidth share)',
    'Internal FIFO: Stores packets as dynamic byte arrays',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.3 Scheduler', level=2)
doc.add_paragraph('The scheduler selects which queue sends the next packet. Three modes:')

doc.add_heading('Strict Priority (SP)', level=3)
doc.add_paragraph(
    'Always selects the queue with the highest prio value that has packets. '
    'Lower-priority queues are only served when all higher-priority queues are empty.'
)

doc.add_heading('Weighted Round Robin (WRR)', level=3)
doc.add_paragraph(
    'Each queue gets bandwidth proportional to its weight. Uses a Deficit Round Robin (DRR) '
    'algorithm with quantum = 1500 bytes. Each queue receives quantum * weight bytes of credit '
    'per round. The queue is served until its credit is depleted, then the next queue is served.'
)
doc.add_paragraph('Example: With weight 3:1, queue 0 gets ~75% bandwidth and queue 1 gets ~25%.')

doc.add_heading('Mixed SP+WRR', level=3)
doc.add_paragraph(
    'Queues with prio > 0 are served first using Strict Priority. '
    'Queues with prio = 0 are served using WRR among themselves.'
)

doc.add_heading('4.4 Port Shaper', level=2)
doc.add_paragraph(
    'The port_shaper applies a token bucket rate limiter at the port level, limiting the '
    'total output bandwidth across all queues.'
)
add_table(
    ['Parameter', 'Default', 'Description'],
    [
        ['port_rate_mbps', '10000.0', 'Maximum port rate in Mbps'],
        ['burst_size_bytes', '16000', 'Maximum burst size in bytes'],
    ]
)

doc.add_heading('4.5 Traffic Monitor', level=2)
doc.add_paragraph(
    'The traffic_monitor collects per-queue statistics and performs windowed rate checking.'
)
add_table(
    ['Metric', 'Description'],
    [
        ['total_bytes', 'Total bytes transmitted'],
        ['total_packets', 'Total packets transmitted'],
        ['actual_rate_mbps', 'Measured transmission rate'],
        ['configured_rate_mbps', 'Expected rate'],
        ['deviation_pct', 'Percentage deviation from expected rate'],
        ['burst_max_bytes', 'Maximum bytes in any window'],
        ['window_violations', 'Windows where rate exceeded tolerance'],
    ]
)
doc.add_paragraph()
doc.add_paragraph(
    'The monitor uses a sliding window (default 10 us) to check whether each queue\'s '
    'instantaneous rate stays within the configured tolerance (default 5%).'
)

doc.add_heading('4.6 Flow Controller', level=2)
doc.add_paragraph(
    'flow_controller is the top-level orchestrator that integrates all components. '
    'Typical workflow:'
)
for i, s in enumerate([
    'Create: flow_controller fc = new();',
    'Configure port: set_port_rate(), set_duration()',
    'Add queues: add_queue(), set_queue_rate(), set_queue_weight()',
    'Set scheduler: set_scheduler_mode()',
    'Push packets: push_packet(queue_id, data)',
    'Start: start_schedule()',
    'Retrieve: while (get_next_scheduled(...)) process packets',
    'Report: report() or get_queue_stats()',
], 1):
    doc.add_paragraph(f'{i}. {s}', style='List Number')

# ====================================================================
#  5. API REFERENCE
# ====================================================================
doc.add_heading('5. API Reference', level=1)

doc.add_heading('5.1 Enumerations', level=2)

doc.add_heading('traffic_model_type_e', level=3)
add_table(
    ['Value', 'Description'],
    [
        ['MODEL_RATE', 'Constant rate model'],
        ['MODEL_TOKEN_BUCKET', 'Token bucket model (rate + burst)'],
        ['MODEL_BURST', 'Burst traffic model'],
        ['MODEL_RANDOM', 'Random distribution model'],
        ['MODEL_STEP', 'Step rate profile model'],
        ['MODEL_CUSTOM', 'User-defined custom model'],
    ]
)

doc.add_paragraph()
doc.add_heading('scheduler_mode_e', level=3)
add_table(
    ['Value', 'Description'],
    [
        ['SCHED_STRICT_PRIORITY', 'Strict Priority scheduling'],
        ['SCHED_WRR', 'Weighted Round Robin scheduling'],
        ['SCHED_SP_WRR_MIXED', 'Mixed SP + WRR scheduling'],
    ]
)

doc.add_paragraph()
doc.add_heading('distribution_e', level=3)
add_table(
    ['Value', 'Description'],
    [
        ['DIST_UNIFORM', 'Uniform distribution'],
        ['DIST_POISSON', 'Poisson (exponential) distribution'],
        ['DIST_NORMAL', 'Normal (Gaussian) distribution'],
    ]
)

doc.add_heading('5.2 Structures', level=2)

doc.add_heading('step_cfg_t', level=3)
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['duration', 'realtime', 'Duration of this step'],
        ['rate_mbps', 'real', 'Rate in Mbps for this step'],
    ]
)

doc.add_paragraph()
doc.add_heading('queue_stats_t', level=3)
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['configured_rate_mbps', 'real', 'Expected rate in Mbps'],
        ['actual_rate_mbps', 'real', 'Measured actual rate in Mbps'],
        ['deviation_pct', 'real', 'Deviation percentage'],
        ['total_bytes', 'longint', 'Total bytes transmitted'],
        ['total_packets', 'longint', 'Total packets transmitted'],
        ['burst_max_bytes', 'longint', 'Maximum bytes in any window'],
        ['window_violations', 'int', 'Number of window violations'],
    ]
)

doc.add_heading('5.3 flow_controller Methods', level=2)

doc.add_heading('Configuration', level=3)
add_table(
    ['Method', 'Parameters', 'Description'],
    [
        ['set_port_rate()', 'real rate_mbps', 'Set total port bandwidth (Mbps)'],
        ['set_port_burst()', 'int burst_bytes', 'Set port burst capacity (bytes)'],
        ['set_duration()', 'realtime dur', 'Set scheduling duration'],
        ['set_time_unit()', 'realtime unit', 'Set time unit (default 1us)'],
        ['set_debug_level()', 'int level', 'Set debug verbosity (0/1/2)'],
        ['set_scheduler_mode()', 'scheduler_mode_e mode', 'Set scheduling algorithm'],
    ]
)

doc.add_paragraph()
doc.add_heading('Queue Management', level=3)
add_table(
    ['Method', 'Parameters', 'Description'],
    [
        ['add_queue()', 'int queue_id, traffic_model_type_e type', 'Create queue with model'],
        ['set_queue_rate()', 'int queue_id, real rate_mbps', 'Set queue rate'],
        ['set_queue_burst()', 'int queue_id, int burst_bytes', 'Set token bucket burst'],
        ['set_queue_priority()', 'int queue_id, int prio', 'Set queue priority (for SP)'],
        ['set_queue_weight()', 'int queue_id, int weight', 'Set queue weight (for WRR)'],
        ['set_burst_param()', 'int qid, int count, realtime pause', 'Configure burst model'],
        ['set_random_param()', 'int qid, real rate, distribution_e dist', 'Configure random model'],
        ['set_step_param()', 'int qid, step_cfg_t steps[$]', 'Configure step model'],
        ['set_queue_model()', 'int qid, traffic_model_base model', 'Set custom model'],
        ['push_packet()', 'int qid, byte unsigned data[$]', 'Enqueue packet'],
    ]
)

doc.add_paragraph()
doc.add_heading('Execution', level=3)
add_table(
    ['Method', 'Parameters', 'Returns', 'Description'],
    [
        ['start_schedule()', 'int pkt_count=-1', 'void', 'Start scheduling (-1=duration mode)'],
        ['get_next_scheduled()', 'out int qid, out byte[] data, out realtime time', 'bit', 'Get next packet (0=done)'],
        ['is_running()', '', 'bit', 'Check if scheduler is active'],
    ]
)

doc.add_paragraph()
doc.add_heading('Monitoring', level=3)
add_table(
    ['Method', 'Parameters', 'Returns', 'Description'],
    [
        ['report()', '', 'void', 'Print statistics report'],
        ['get_queue_stats()', 'int queue_id', 'queue_stats_t', 'Get statistics for queue'],
        ['monitor_packet()', 'int qid, byte[] data, realtime time', 'void', 'Record received packet'],
    ]
)

# ====================================================================
#  6. USAGE EXAMPLES
# ====================================================================
doc.add_heading('6. Usage Examples', level=1)

doc.add_heading('6.1 Basic Rate-Limited Traffic', level=2)
doc.add_paragraph('Send constant-rate traffic from a single queue at 1 Gbps through a 10 Gbps port.')
add_code('''flow_controller fc = new();
byte unsigned pkt[$];
int qid;
byte unsigned out_data[$];
realtime send_time;

// Configure port and duration
fc.set_port_rate(10000.0);              // 10 Gbps port
fc.set_duration(100us);

// Add one queue at 1 Gbps
fc.add_queue(0, MODEL_RATE);
fc.set_queue_rate(0, 1000.0);           // 1 Gbps

// Generate 100 packets of 1500 bytes
pkt = {};
for (int i = 0; i < 1500; i++) pkt.push_back(i % 256);
for (int i = 0; i < 100; i++)
    fc.push_packet(0, pkt);

// Run
fc.start_schedule();
while (fc.get_next_scheduled(qid, out_data, send_time)) begin
    $display("[%0t] Sent %0d bytes from queue %0d", send_time, out_data.size(), qid);
end
fc.report();''')

doc.add_heading('6.2 Token Bucket with Burst', level=2)
doc.add_paragraph(
    'Token bucket allows burst transmission up to burst size, then rate-limits.'
)
add_code('''fc.add_queue(0, MODEL_TOKEN_BUCKET);
fc.set_queue_rate(0, 500.0);            // 500 Mbps sustained rate
fc.set_queue_burst(0, 16000);           // 16 KB burst capacity

// With 1000-byte packets: first 16 packets arrive instantly (burst),
// then packets are rate-limited to 500 Mbps''')

doc.add_heading('6.3 Multi-Queue WRR Scheduling', level=2)
doc.add_paragraph('Weighted Round Robin distributes bandwidth proportionally to queue weights.')
add_code('''flow_controller fc = new();
byte unsigned pkt[$];

fc.set_port_rate(10000.0);
fc.set_duration(100us);

// Add 3 queues with weights 4:2:1
fc.add_queue(0, MODEL_RATE);
fc.set_queue_rate(0, 5000.0);
fc.set_queue_weight(0, 4);              // 4/(4+2+1) = 57% bandwidth

fc.add_queue(1, MODEL_RATE);
fc.set_queue_rate(1, 5000.0);
fc.set_queue_weight(1, 2);              // 2/(4+2+1) = 29% bandwidth

fc.add_queue(2, MODEL_RATE);
fc.set_queue_rate(2, 5000.0);
fc.set_queue_weight(2, 1);              // 1/(4+2+1) = 14% bandwidth

fc.set_scheduler_mode(SCHED_WRR);

// Push packets to all queues
pkt = {};
for (int i = 0; i < 1500; i++) pkt.push_back(8'hAA);
for (int i = 0; i < 500; i++) begin
    fc.push_packet(0, pkt);
    fc.push_packet(1, pkt);
    fc.push_packet(2, pkt);
end

fc.start_schedule();
// ... get_next_scheduled loop ...
fc.report();''')

doc.add_heading('6.4 Strict Priority Scheduling', level=2)
doc.add_paragraph('Higher-priority queues are always served first.')
add_code('''fc.add_queue(0, MODEL_RATE);
fc.set_queue_rate(0, 5000.0);
fc.set_queue_priority(0, 3);            // Highest priority

fc.add_queue(1, MODEL_RATE);
fc.set_queue_rate(1, 3000.0);
fc.set_queue_priority(1, 2);            // Medium priority

fc.add_queue(2, MODEL_RATE);
fc.set_queue_rate(2, 1000.0);
fc.set_queue_priority(2, 1);            // Lowest priority

fc.set_scheduler_mode(SCHED_STRICT_PRIORITY);
// Queue 0 served exclusively until empty, then queue 1, then queue 2.''')

doc.add_heading('6.5 Mixed SP+WRR Scheduling', level=2)
doc.add_paragraph('Combine SP for critical traffic with WRR for best-effort traffic.')
add_code('''// High-priority control plane traffic (SP)
fc.add_queue(0, MODEL_RATE);
fc.set_queue_rate(0, 1000.0);
fc.set_queue_priority(0, 5);            // prio > 0: SP mode

// Best-effort data plane traffic (WRR, prio = 0)
fc.add_queue(1, MODEL_RATE);
fc.set_queue_rate(1, 3000.0);
fc.set_queue_priority(1, 0);            // prio = 0: WRR mode
fc.set_queue_weight(1, 3);

fc.add_queue(2, MODEL_RATE);
fc.set_queue_rate(2, 3000.0);
fc.set_queue_priority(2, 0);
fc.set_queue_weight(2, 1);

fc.set_scheduler_mode(SCHED_SP_WRR_MIXED);
// Queue 0 always served first (SP).
// When queue 0 is empty, queues 1 and 2 share bandwidth at 3:1 ratio (WRR).''')

doc.add_heading('6.6 Burst Traffic Pattern', level=2)
doc.add_paragraph('Send N packets back-to-back, then pause, then repeat.')
add_code('''fc.add_queue(0, MODEL_BURST);
fc.set_burst_param(0,
    5,              // send_count: 5 packets per burst
    10us            // pause_time: 10 us between bursts
);

// Traffic pattern: [pkt pkt pkt pkt pkt] --10us-- [pkt pkt pkt pkt pkt] --10us-- ...''')

doc.add_heading('6.7 Random Traffic Model', level=2)
doc.add_paragraph('Generate traffic with stochastic inter-packet gaps.')
add_code('''// Uniform random
fc.add_queue(0, MODEL_RANDOM);
fc.set_random_param(0, 2000.0, DIST_UNIFORM);    // 2 Gbps avg, uniform

// Poisson process (memoryless)
fc.add_queue(1, MODEL_RANDOM);
fc.set_random_param(1, 1000.0, DIST_POISSON);    // 1 Gbps avg, poisson

// Gaussian random
fc.add_queue(2, MODEL_RANDOM);
fc.set_random_param(2, 500.0, DIST_NORMAL);      // 500 Mbps avg, normal''')

doc.add_heading('6.8 Step Rate Profile', level=2)
doc.add_paragraph('Define a multi-segment rate profile that changes over time.')
add_code('''fc.add_queue(0, MODEL_STEP);

// Define rate steps
begin
    step_cfg_t steps[$];
    step_cfg_t s;

    s.duration = 10us;  s.rate_mbps = 1000.0;   // 0-10us:  1 Gbps
    steps.push_back(s);
    s.duration = 20us;  s.rate_mbps = 5000.0;   // 10-30us: 5 Gbps
    steps.push_back(s);
    s.duration = 10us;  s.rate_mbps = 500.0;    // 30-40us: 500 Mbps
    steps.push_back(s);

    fc.set_step_param(0, steps);
end

// To loop the profile continuously:
// step_model sm;
// $cast(sm, queues[0].model);
// sm.loop_enable = 1;''')

doc.add_heading('6.9 Custom Traffic Model', level=2)
doc.add_paragraph('Create your own traffic model by extending traffic_model_base.')
add_code('''// Define custom model
class my_model extends traffic_model_base;
    real base_rate_mbps;
    int pkt_count;

    function new();
        super.new();
        model_type = MODEL_CUSTOM;
        base_rate_mbps = 1000.0;
        pkt_count = 0;
    endfunction

    virtual function realtime get_interval(int pkt_size, realtime current_time);
        real rate;
        pkt_count++;
        // Increase rate every 10 packets
        rate = base_rate_mbps * (1 + pkt_count / 10);
        return (pkt_size * 8.0) / (rate * 1_000_000.0) * 1s;
    endfunction

    virtual function void reset();
        pkt_count = 0;
    endfunction

    virtual function string to_string();
        return $sformatf("my_model: base_rate=%.1f Mbps", base_rate_mbps);
    endfunction
endclass

// Use custom model
fc.add_queue(0, MODEL_CUSTOM);
begin
    my_model m = new();
    m.base_rate_mbps = 500.0;
    fc.set_queue_model(0, m);
end''')

doc.add_heading('6.10 Statistics & Monitoring', level=2)
doc.add_paragraph('Access detailed per-queue statistics after simulation.')
add_code('''// After scheduling completes:
queue_stats_t stats;

// Method 1: Print formatted report
fc.report();
// Output:
// === Flow Controller Report ===
// Duration: 100.000 us | Window: 10.000 us | Tolerance: 5.0%
// Queue | Config Rate | Actual Rate | Dev    | Packets | Bytes     | Violations
// 0     | 1000.0 Mbps | 1090.9 Mbps |  9.1%  |      12 |     12000 | 0
// Total: 12 packets, 12000 bytes
// === END ===

// Method 2: Access stats programmatically
stats = fc.get_queue_stats(0);
$display("Queue 0: rate=%.1f Mbps, deviation=%.1f%%, packets=%0d",
         stats.actual_rate_mbps, stats.deviation_pct, stats.total_packets);

// Check rate accuracy
if (stats.deviation_pct > 10.0 || stats.deviation_pct < -10.0)
    $error("Rate deviation too high: %.1f%%", stats.deviation_pct);

// Check for window violations
if (stats.window_violations > 0)
    $warning("Queue 0 had %0d window violations", stats.window_violations);''')

doc.add_heading('6.11 Complete Test Bench Example', level=2)
doc.add_paragraph('A complete test bench with multi-queue WRR scheduling and statistics verification.')
add_code('''program complete_testbench;
    `include "core/flow_controller.sv"

    initial begin
        flow_controller fc = new();
        byte unsigned pkt[$];
        int qid;
        byte unsigned out_data[$];
        realtime send_time;
        int q0_count, q1_count, total;
        queue_stats_t s0, s1;

        // ---- Configuration ----
        fc.set_port_rate(10000.0);          // 10 Gbps port
        fc.set_duration(200us);             // 200 us simulation
        fc.set_debug_level(0);              // 0=quiet, 2=trace

        // ---- Queue Setup ----
        fc.add_queue(0, MODEL_RATE);
        fc.set_queue_rate(0, 3000.0);       // 3 Gbps
        fc.set_queue_weight(0, 3);

        fc.add_queue(1, MODEL_RATE);
        fc.set_queue_rate(1, 2000.0);       // 2 Gbps
        fc.set_queue_weight(1, 1);

        fc.set_scheduler_mode(SCHED_WRR);

        // ---- Packet Generation ----
        pkt = {};
        for (int i = 0; i < 1500; i++) pkt.push_back(i % 256);
        for (int i = 0; i < 1000; i++) begin
            fc.push_packet(0, pkt);
            fc.push_packet(1, pkt);
        end

        // ---- Scheduling ----
        q0_count = 0;
        q1_count = 0;
        total = 0;
        fc.start_schedule();

        while (fc.get_next_scheduled(qid, out_data, send_time)) begin
            if (qid == 0) q0_count++;
            else q1_count++;
            total++;
        end

        // ---- Results ----
        fc.report();

        s0 = fc.get_queue_stats(0);
        s1 = fc.get_queue_stats(1);

        $display("");
        $display("Scheduling Results:");
        $display("  Total packets sent: %0d", total);
        $display("  Queue 0: %0d packets (%.1f%%)", q0_count, 100.0*q0_count/total);
        $display("  Queue 1: %0d packets (%.1f%%)", q1_count, 100.0*q1_count/total);
        $display("");
        $display("Rate Verification:");
        $display("  Queue 0: configured=%.1f, actual=%.1f, dev=%.1f%%",
                 s0.configured_rate_mbps, s0.actual_rate_mbps, s0.deviation_pct);
        $display("  Queue 1: configured=%.1f, actual=%.1f, dev=%.1f%%",
                 s1.configured_rate_mbps, s1.actual_rate_mbps, s1.deviation_pct);

        // ---- Assertions ----
        if (q0_count <= q1_count)
            $error("WRR failed: q0 (weight=3) should send more than q1 (weight=1)");
        if (s0.total_bytes != q0_count * 1500)
            $error("Stats mismatch: bytes != packets * pkt_size");

        $display("\\nTest completed successfully.");
    end
endprogram''')

# ====================================================================
#  7. TEST SUITE
# ====================================================================
doc.add_heading('7. Test Suite', level=1)
doc.add_paragraph('The project includes 89 test cases across 6 test files:')

add_table(
    ['Test File', 'Tests', 'Coverage'],
    [
        ['test_traffic_models.sv', '19', 'All 5 models: rate, token bucket, burst, random, step'],
        ['test_queue.sv', '12', 'FIFO: push/pop, depth, flush, model assignment, peek'],
        ['test_scheduler.sv', '17', 'SP, WRR (3:1, 1:1, 4:2:1), drain, mixed, preemption, reset'],
        ['test_port_shaper.sv', '4', 'Token refill, burst exhaustion, wait time'],
        ['test_flow_controller.sv', '29', 'Integration: integrity, SP/WRR, rate, stats, burst, duration, monotonicity'],
        ['test_monitor.sv', '8', 'Statistics, multi-queue tracking, deviation, debug'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('Run all tests:')
add_code('make test_all')
doc.add_paragraph('Run individual test:')
add_code('make test_scheduler    # or test_queue, test_traffic_models, etc.')

# ====================================================================
#  8. BUILDING & RUNNING
# ====================================================================
doc.add_heading('8. Building & Running', level=1)

doc.add_heading('Prerequisites', level=2)
doc.add_paragraph('Synopsys VCS simulator with valid license.')

doc.add_heading('Environment Setup', level=2)
add_code('''export VCS_HOME=/opt/synopsys/vcs/Q-2020.03-SP2-7
export PATH=$VCS_HOME/bin:$PATH
export LM_LICENSE_FILE=30000@<license_server>''')

doc.add_heading('Makefile Targets', level=2)
add_table(
    ['Target', 'Description'],
    [
        ['make test_all', 'Run all 6 test suites'],
        ['make test_traffic_models', 'Test traffic models'],
        ['make test_queue', 'Test queue operations'],
        ['make test_scheduler', 'Test scheduling algorithms'],
        ['make test_port_shaper', 'Test port rate limiter'],
        ['make test_flow_controller', 'Test flow controller integration'],
        ['make test_monitor', 'Test statistics monitor'],
        ['make clean', 'Remove build artifacts'],
    ]
)

doc.add_paragraph()
doc.add_heading('Compilation Flags', level=2)
add_code('VCS_FLAGS := -full64 -sverilog -timescale=1ns/1ps -f filelist.f +incdir+src')

doc.add_heading('Custom Test Bench', level=2)
doc.add_paragraph('To compile and run your own test bench:')
add_code('''vcs -full64 -sverilog -timescale=1ns/1ps \\
    -f filelist.f +incdir+src \\
    your_testbench.sv -o simv
./simv''')

doc.add_heading('Debug Mode', level=2)
doc.add_paragraph('Enable packet-level tracing:')
add_code('''fc.set_debug_level(2);
// Output: [8000] TX queue=0 pkt_size=1000
//         [16000] TX queue=0 pkt_size=1000''')

# ── Save ──
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        'Flow_Control_User_Manual.docx')
doc.save(out_path)
print(f'Document saved to: {out_path}')
